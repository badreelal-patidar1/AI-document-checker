import os
import io
import fitz
import pytest
import asyncio
from docx import Document
from docx.shared import RGBColor
from fastapi import UploadFile
from utils import (
    validate_file,
    extract_text_from_pdf,
    extract_text_from_docx,
    convert_custom_text_to_html,
    save_pdf_text,
    save_docx_text,
    chunk_text,
    ensure_dirs,
    _hex_to_rgbcolor,
    _add_runs_from_text,
    _process_block_into_para
)

# ---------- SETUP FIXTURES ----------

@pytest.fixture
def tmp_pdf(tmp_path):
    """Create a small sample PDF for testing."""
    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Hello PDF Bold", fontname="helvetica-bold", fontsize=12)
    doc.save(path)
    return str(path)

@pytest.fixture
def tmp_docx(tmp_path):
    """Create a small DOCX for testing."""
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Heading 1", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Hello Bold")
    run.bold = True
    doc.save(path)
    return str(path)

# ---------- VALIDATE FILE ----------

@pytest.mark.asyncio
async def test_validate_file_pdf(tmp_pdf):
    with open(tmp_pdf, "rb") as f:
        upload = UploadFile(filename="test.pdf", file=io.BytesIO(f.read()))
    valid, msg = await validate_file(upload)
    assert valid is True
    assert "Allowed" in msg

@pytest.mark.asyncio
async def test_validate_file_invalid_extension():
    fake_file = UploadFile(filename="invalid.txt", file=io.BytesIO(b"data"))
    valid, msg = await validate_file(fake_file)
    assert valid is False
    assert "PDF or DOCX" in msg

# ---------- PDF TEXT EXTRACTION ----------

@pytest.mark.asyncio
async def test_extract_text_from_pdf(tmp_pdf):
    text = await extract_text_from_pdf(tmp_pdf)
    assert "Hello" in text
    assert "<b>" in text  # should detect bold font

# ---------- DOCX TEXT EXTRACTION ----------

@pytest.mark.asyncio
async def test_extract_text_from_docx(tmp_docx):
    text = await extract_text_from_docx(tmp_docx)
    assert "Hello" in text
    assert "**Hello Bold**" in text or "*Hello Bold*" in text

# ---------- HTML CONVERSION ----------

@pytest.mark.asyncio
async def test_convert_custom_text_to_html():
    text = "<#FF0000>Hello</#FF0000>\n\nWorld"
    html = await convert_custom_text_to_html(text)
    assert "<span" in html
    assert "color:#FF0000" in html
    assert "<p>" in html

# ---------- PDF SAVE ----------

@pytest.mark.asyncio
async def test_save_pdf_text(tmp_path):
    text = "<#0000FF>Hello PDF</#0000FF>"
    pdf_path = tmp_path / "output.pdf"
    out_path = await save_pdf_text(text, str(pdf_path))
    assert os.path.exists(out_path)
    assert out_path.endswith(".pdf")

# ---------- DOCX SAVE + COLOR + MARKDOWN ----------

@pytest.mark.asyncio
async def test_save_docx_text_and_open(tmp_path):
    text = """
# Title
## Subtitle

• Bullet 1
1. Numbered item

Normal **bold** and *italic* and __underline__

<#FF0000>Red text</#FF0000>
"""
    path = tmp_path / "out.docx"
    await save_docx_text(str(path), text)
    assert os.path.exists(path)

    doc = Document(path)
    all_text = " ".join([p.text for p in doc.paragraphs])
    assert "Title" in all_text
    assert "bold" in all_text
    assert "Red text" in all_text

# ---------- HEX TO RGB ----------

@pytest.mark.asyncio
async def test_hex_to_rgbcolor_valid():
    color = await _hex_to_rgbcolor("FF0000")
    assert color == RGBColor(255, 0, 0)  

@pytest.mark.asyncio
async def test_hex_to_rgbcolor_invalid():
    with pytest.raises(ValueError):
        await _hex_to_rgbcolor("GG0000")

# ---------- ADD RUNS ----------

@pytest.mark.asyncio
async def test_add_runs_from_text(tmp_path):
    doc = Document()
    para = doc.add_paragraph()
    await _add_runs_from_text(para, "Hello **Bold** __Under__ *Italics*")
    runs = [r.text for r in para.runs]
    assert any("Hello" in r for r in runs)
    assert any(r.bold for r in para.runs if "Bold" in r.text)
    assert any(r.underline for r in para.runs if "Under" in r.text)
    assert any(r.italic for r in para.runs if "Italics" in r.text)

# ---------- PROCESS BLOCK ----------

@pytest.mark.asyncio
async def test_process_block_into_para_heading():
    doc = Document()
    para = await _process_block_into_para(doc, "# Heading 1")
    assert para.style.name.startswith("Heading")

@pytest.mark.asyncio
async def test_process_block_into_para_colored():
    doc = Document()
    para = await _process_block_into_para(doc, "Normal <#FF0000>Red</#FF0000> Text")
    text = " ".join([r.text for r in para.runs])
    assert "Red" in text

# ---------- CHUNK TEXT ----------

@pytest.mark.asyncio
async def test_chunk_text():
    text = "A" * 15000
    chunks = await chunk_text(text, 7000)
    assert len(chunks) == 3
    assert sum(len(c) for c in chunks) == len(text)

# ---------- ENSURE DIRS ----------

def test_ensure_dirs(tmp_path):
    dirs = [tmp_path / "a", tmp_path / "b"]
    ensure_dirs(dirs)
    for d in dirs:
        assert os.path.exists(d)
