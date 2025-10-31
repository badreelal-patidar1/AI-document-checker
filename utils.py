# utils.py
import fitz  # PyMuPDF
import traceback
import os
import re
import magic
import asyncio
from typing import List
from fastapi import UploadFile
from docx import Document
from docx.shared import RGBColor
from xhtml2pdf import pisa

CHUNK_SIZE = 7000  # characters per chunk for LLM prompts (adjust as needed)

_hex_color_tag_re = re.compile(r"<#([0-9A-Fa-f]{6})>(.*?)</#\1>", re.DOTALL)

async def validate_file(file: UploadFile) -> tuple[bool, str]:
    """
    Validate the uploaded file to ensure it is either a PDF or DOCX and that its MIME type matches
    allowed document formats.

    Args:
        file (UploadFile): The file uploaded via FastAPI.

    Returns:
        tuple[bool, str]: A tuple where the first element is True if the file is valid, False otherwise.
                          The second element is an explanatory message.

    The function checks the file extension and MIME type against allowed types:
      - application/pdf
      - application/msword
      - application/vnd.openxmlformats-officedocument.wordprocessingml.document

    It reads a small chunk from the file for MIME detection and resets the file pointer afterward.
    """
    file_name = file.filename.lower()
    if (file_name.endswith(".pdf") or file_name.endswith(".docx")):
        # Read small chunk for MIME detection
        head = await file.read(2048)
        await file.seek(0)  # reset file pointer for future reads
        mime_type = magic.from_buffer(head, mime=True)
        allowed_types = ["application/pdf", "application/msword",
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
        
        if mime_type in allowed_types:
            return True, "Allowed file"
        else:
            msg = f"Not allowed: {file_name} ({mime_type})"
            return False, msg
    else:
        msg = f"File must be a PDF or DOCX"
        return False, msg

async def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file at the given file path, preserving structure and style information.

    This function reads a PDF document and attempts to preserve various text attributes, including:
        - Structure (paragraph/line grouping)
        - Font family and font size
        - Bold and italic styling
        - Font color (as #RRGGBB HEX codes)

    The extracted and formatted text uses readable custom tags such as <b>, <i>, <size=12>, <font=Helvetica>, and <#RRGGBB>.

    Args:
        file_path (str): The path to the PDF file.

    Returns:
        str: The extracted and formatted text from the PDF.
    """
    
    formatted_text = []

    with fitz.open(file_path) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            blocks = (await asyncio.to_thread(page.get_text, "dict"))["blocks"]
            for block in blocks:
                if "lines" not in block:
                    continue

                paragraph = []
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if not text:
                            continue

                        font_name = span.get("font", "")
                        font_size = span.get("size", 12)
                        is_bold = "Bold" in font_name
                        is_italic = "Italic" in font_name or "Oblique" in font_name

                        # Extract base font family (before '-' or ',' etc.)
                        font_family = font_name.split("-")[0] if "-" in font_name else font_name

                        # Convert color int → RGB → HEX
                        color_val = span.get("color", 0)
                        r, g, b = (color_val >> 16) & 255, (color_val >> 8) & 255, color_val & 255
                        color_hex = f"#{r:02X}{g:02X}{b:02X}"

                        # Apply formatting tags (nest carefully)
                        styled = f"<size={font_size:.1f}><font={font_family}><{color_hex}>{text}</{color_hex}></font></size>"
                        if is_bold:
                            styled = f"<b>{styled}</b>"
                        if is_italic:
                            styled = f"<i>{styled}</i>"

                        paragraph.append(styled)

                if paragraph:
                    formatted_text.append("".join(paragraph))
                    formatted_text.append("\n\n")  # paragraph spacing

    return "".join(formatted_text).strip()

async def convert_custom_text_to_html(text: str) -> str:
    """
    Convert a custom-formatted text string (with tags for color, bold, italic, etc.)
    into valid HTML with inline CSS styles.

    This function is intended to parse text containing special tags,
    such as <#RRGGBB> for color, <b> for bold, and <i> for italics,
    and output a corresponding HTML document where these styles are
    preserved using standard HTML elements and inline CSS.

    Args:
        text (str): The input string with custom formatting tags.

    Returns:
        str: A complete HTML document as a string with equivalent formatting.
    """
   
    html = text

    # Replace all opening color tags like <#8B0000>
    html = re.sub(r"<#([A-Fa-f0-9]{6})>", r'<span style="color:#\1;">', html)

    # Replace all closing color tags like </#8B0000>
    html = re.sub(r"</#([A-Fa-f0-9]{6})>", "</span>", html)

    # Make sure line breaks become paragraphs for readability
    html = html.replace("\n\n", "</p><p>").replace("\n", "<br>")

    # Wrap final HTML
    html = f"""
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.5;
                font-size: 12pt;
                margin: 20px;
            }}
            b {{ font-weight: bold; }}
            i {{ font-style: italic; }}
            p {{ margin-bottom: 10px; }}
        </style>
    </head>
    <body><p>{html}</p></body>
    </html>
    """
    return html

async def save_pdf_text(text: str, pdf_path: str):
    """
    Convert custom-formatted text to a PDF file.

    This function takes a string `text` containing custom formatting tags
    (such as color, bold, and italic), converts it to HTML using
    `convert_custom_text_to_html`, and then writes the content to a PDF file
    at the specified `pdf_path` using the xhtml2pdf library.

    Args:
        text (str): The text content to convert, possibly with custom tags.
        pdf_path (str): The file path where the PDF should be saved.

    Returns:
        str: The path to the saved PDF file.
    """
    html = await convert_custom_text_to_html(text)
    with open(pdf_path, "wb") as pdf_file:
        await asyncio.to_thread(pisa.CreatePDF, html, dest=pdf_file)
    return pdf_path

async def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file while preserving structure and formatting.
    Keeps headings, lists, tables, and inline styles like bold, italic,
    underline, and text color.

    Args:
        file_path (str): The path to the DOCX file.

    Returns:
        str: The extracted and formatted text from the DOCX.
    """
    async def get_color_tag(run):
        """
        Convert RGBColor to readable color tags like <red>text</red>

        Args:
            run (docx.text.run.Run): The run object to convert.

        Returns:
            str: The color tag or None if no color is detected. If a color is detected, it returns a readable color tag like <red>text</red>.
        """
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            r, g, b = rgb[0], rgb[1], rgb[2]
            # Simple color detection for common colors
            if r > 200 and g < 80 and b < 80:
                return "red"
            elif b > 200 and r < 80 and g < 80:
                return "blue"
            elif g > 150 and r < 100:
                return "green"
            elif r > 200 and g > 200 and b < 100:
                return "yellow"
            elif r > 100 and g > 100 and b > 100:
                return "gray"
            else:
                return f"#{rgb}"
        return None

    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        style_name = getattr(para.style, "name", "").lower()

        # Detect headings
        if style_name.startswith("heading"):
            level = ''.join(filter(str.isdigit, style_name)) or '1'
            lines.append(f"{'#' * int(level)} {para.text.strip()}")
            continue

        # Bulleted list
        elif "list bullet" in style_name or "list paragraph" in style_name:
            prefix = "• "
        # Numbered list
        elif "list number" in style_name:
            prefix = "1. "
        else:
            prefix = ""

        formatted_text = ""
        for run in para.runs:
            text = run.text
            if not text:
                continue

            # Detect color
            color_tag = await get_color_tag(run)

            # Apply bold/italic/underline
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"__{text}__"

            # Wrap in color tag if applicable
            if color_tag:
                text = f"<{color_tag}>{text}</{color_tag}>"

            formatted_text += text

        lines.append(prefix + formatted_text.strip())

    # Handle tables (preserve structure)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(row_text))
        lines.append("")  # blank line after table

    return "\n\n".join(lines).strip()


async def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> List[str]:
    """
    Split a given string into smaller chunks of at most `chunk_size` characters each.

    Args:
        text (str): The input text to chunk.
        chunk_size (int, optional): Maximum size of each chunk. Defaults to CHUNK_SIZE.

    Returns:
        List[str]: A list of string chunks, where each chunk is at most `chunk_size` characters long.
    """
    
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunks.append(text[start:end])
        start = end
    return chunks

async def _hex_to_rgbcolor(hexstr: str):
    """
    Convert a 6-character hex color string (e.g., '365F91') to a docx.shared.RGBColor object.

    Args:
        hexstr (str): The input hex color string to convert.

    Returns:
        docx.shared.RGBColor: The RGBColor object representing the color.
    """
    hexstr = hexstr.strip()
    if len(hexstr) != 6:
        raise ValueError("Hex color must be 6 hex digits")
    r = int(hexstr[0:2], 16)
    g = int(hexstr[2:4], 16)
    b = int(hexstr[4:6], 16)
    return RGBColor(r, g, b)

async def _add_runs_from_text(para, text: str, color_rgb=None):
    """
    Add runs to `para` from `text`, interpreting **bold**, __underline__, *italic*.
    color_rgb: docx.shared.RGBColor or None — applied to all created runs.

    Args:
        para (docx.text.paragraph.Paragraph): The paragraph object to add runs to.
        text (str): The text to add runs to.
        color_rgb (docx.shared.RGBColor, optional): The color to apply to all created runs. Defaults to None.

    Returns:
        None
    """
    if not text:
        return

    # Patterns for inline formats
    # We will search for earliest occurrence among these markers.
    bold_re = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)
    underline_re = re.compile(r"__(.+?)__", re.DOTALL)
    italic_re = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", re.DOTALL)  # single * not part of **
    # Note: italic regex tries to avoid matching the ** markers.

    i = 0
    while text:
        # Find earliest match among bold, underline, italic
        m_b = bold_re.search(text)
        m_u = underline_re.search(text)
        m_i = italic_re.search(text)

        # pick earliest found (lowest start index), ignore None
        matches = [(m_b, "b"), (m_u, "u"), (m_i, "i")]
        matches = [(m, t) for m, t in matches if m]
        if not matches:
            # no more formatting markers; add remaining text as a plain run
            run = para.add_run(text)
            if color_rgb:
                run.font.color.rgb = color_rgb
            break

        # choose earliest match
        m, kind = min(matches, key=lambda mt: mt[0].start())

        # text before match -> plain run
        if m.start() > 0:
            before = text[: m.start()]
            run_before = para.add_run(before)
            if color_rgb:
                run_before.font.color.rgb = color_rgb

        # formatted content
        content = m.group(1)
        run_fmt = para.add_run(content)
        if kind == "b":
            run_fmt.bold = True
        elif kind == "u":
            run_fmt.underline = True
        elif kind == "i":
            run_fmt.italic = True
        if color_rgb:
            run_fmt.font.color.rgb = color_rgb

        # continue after this match
        text = text[m.end() :]

async def _process_block_into_para(doc, block: str):
    """
    Create either a heading, list paragraph, or normal paragraph and add runs with formatting + color support.

    Args:
        doc (docx.document.Document): The document object to add the paragraph to.
        block (str): The text block to process.

    Returns:
        docx.text.paragraph.Paragraph: The created paragraph object.
    """
    # Headings
    if block.startswith("#"):
        level = len(block) - len(block.lstrip("#"))
        heading_text = block.lstrip("#").strip()
        return doc.add_heading(heading_text, level=level)

    # Bullet list
    if block.startswith("• "):
        para = doc.add_paragraph("", style="List Bullet")
        block = block[2:].strip()
    # Numbered list
    elif block.startswith("1. "):
        para = doc.add_paragraph("", style="List Number")
        block = block[3:].strip()
    else:
        para = doc.add_paragraph("")

    # Now handle color-tags by splitting into segments: plain & color-tagged parts
    parts = _hex_color_tag_re.split(block)
    # re.split with capture groups returns sequence: [prefix, color1, inner1, suffix, ...]
    # We need to iterate. But simpler: use finditer to iterate in order.

    pos = 0
    for m in _hex_color_tag_re.finditer(block):
        # text before color tag
        if m.start() > pos:
            prefix = block[pos : m.start()]
            await _add_runs_from_text(para, prefix, color_rgb=None)
        # colored inner text
        color_hex = m.group(1)
        inner = m.group(2)
        try:
            rgb = await _hex_to_rgbcolor(color_hex)
        except Exception:
            rgb = None
        await _add_runs_from_text(para, inner, color_rgb=rgb)
        pos = m.end()

    # remaining suffix after last color tag
    if pos < len(block):
        suffix = block[pos:]
        await _add_runs_from_text(para, suffix, color_rgb=None)

    return para

async def save_docx_text(path: str, text: str):
    """
    Save structured text (Markdown-like + color tags + inline formatting) into a .docx file.
    Supports:
      - Headings (#, ##, ###)
      - Bullet & numbered lists
      - Bold (**text**), Italic (*text*), Underline (__text__)
      - Hex color tags like <#FF0000>red text</#FF0000>

    Args:
        path (str): The path to the DOCX file to save.
        text (str): The text to save to the DOCX file.

    Returns:
        None
    """
    doc = Document()

    # Split blocks on double newline; keep order and structure
    for block in re.split(r"\n\s*\n", text.strip()):
        block = block.strip()
        if not block:
            continue
        await _process_block_into_para(doc, block)

    doc.save(path)

def ensure_dirs(dir_list):
    """
    Ensure directories exist.
    Creates directories if they do not exist.
    Does not raise an error if the directories already exist.
    Args:
        dir_list (list): The list of directories to ensure exist.

    Returns:
        None
    """
    for d in dir_list:
        os.makedirs(d, exist_ok=True)
